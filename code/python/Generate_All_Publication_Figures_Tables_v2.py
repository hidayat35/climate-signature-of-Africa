"""
Generate_All_Publication_Figures_Tables_v2.py

PATCH OF v1 (which crashed on FigureB1 column 'pathway')

What changed
------------
The Module B CSV files don't all use the column names I guessed in v1.
v2 fixes this by:
  1. Reading each CSV first and PRINTING all column names (so any future
     mismatch is immediately visible in the console).
  2. Discovering column names dynamically using fuzzy keyword matching,
     not hardcoded literal names.
  3. Falling back gracefully if a column isn't found — the figure will
     still be produced where possible, with a console warning instead
     of a crash.

Everything else (Figures 1, 2, 3, S1, S2, all tables) is identical.

Run
---
python Generate_All_Publication_Figures_Tables_v2.py

The output folder is the same as before. Existing files in it are overwritten.
"""

import os
import sys
import shutil
from pathlib import Path
import warnings

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl
from matplotlib.colors import LinearSegmentedColormap, ListedColormap, BoundaryNorm
from matplotlib.patches import Patch

try:
    import rasterio
    HAS_RASTERIO = True
except ImportError:
    HAS_RASTERIO = False

try:
    import geopandas as gpd
    HAS_GEOPANDAS = True
except ImportError:
    HAS_GEOPANDAS = False

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

warnings.filterwarnings('ignore')

# ============================================================
# CONFIGURATION
# ============================================================
DATA_ROOT       = Path(r"D:\Claude idea\PhD_Paper3_Data")
SPEI_OUTPUT_DIR = DATA_ROOT / "step5_spei_output"
FIGS_DIR        = SPEI_OUTPUT_DIR / "figures_for_paper"
MODULEB_DIR     = FIGS_DIR / "ModuleB"
STATS_DIR       = SPEI_OUTPUT_DIR / "statistics_for_paper"
OUTPUT_DIR      = FIGS_DIR / "clean figs and table"
SHAPEFILE       = Path(r"D:\Claude idea\ipc_africa_5_regions.shp")

DPI = 1000

REGION_ORDER  = ['MED', 'SAH', 'WAF', 'EAF', 'SAF']
REGION_NAMES  = {'MED': 'Mediterranean',  'SAH': 'Sahara-Sahel',
                 'WAF': 'West Africa',    'EAF': 'East Africa',
                 'SAF': 'Southern Africa'}
REGION_COLORS = {'MED': '#d62728', 'SAH': '#ff7f0e',
                 'WAF': '#2ca02c', 'EAF': '#1f77b4', 'SAF': '#9467bd'}

mpl.rcParams.update({
    'font.family': 'DejaVu Sans', 'font.size': 10,
    'axes.titlesize': 11, 'axes.labelsize': 10,
    'xtick.labelsize': 9, 'ytick.labelsize': 9,
    'legend.fontsize': 9, 'figure.titlesize': 12,
    'savefig.bbox': 'tight', 'savefig.facecolor': 'white',
    'axes.spines.top': False, 'axes.spines.right': False,
})


# ============================================================
# DYNAMIC COLUMN DISCOVERY
# ============================================================
def find_col(df, keywords, label='?'):
    """
    Find a column name in df by trying each keyword (substring match,
    case-insensitive). Returns the first matching column name or None.

    keywords: list of substrings to look for. First match wins.
    """
    for kw in keywords:
        for c in df.columns:
            if kw.lower() in c.lower():
                return c
    print(f"    [warn] Could not find column for '{label}' "
          f"(tried {keywords}). Columns present: {list(df.columns)}")
    return None


def print_columns(name, df):
    """Print discovered columns for transparency."""
    print(f"    Columns in {name}: {list(df.columns)}")


# ============================================================
# LOGGING
# ============================================================
class Log:
    def __init__(self):
        self.entries = []
        self.n_ok = self.n_skip = self.n_fail = 0
    def ok(self, msg):    print(f"  ✓ {msg}"); self.entries.append(('OK', msg));   self.n_ok += 1
    def skip(self, msg):  print(f"  ⊘ {msg}"); self.entries.append(('SKIP', msg)); self.n_skip += 1
    def fail(self, msg):  print(f"  ✗ {msg}"); self.entries.append(('FAIL', msg)); self.n_fail += 1
    def info(self, msg):  print(f"    {msg}")
    def section(self, t): print(f"\n{'='*70}\n{t}\n{'='*70}")
    def summary(self):
        print(f"\n{'='*70}\nSUMMARY: {self.n_ok} OK, {self.n_skip} skipped, "
              f"{self.n_fail} failed\nOutput: {OUTPUT_DIR}\n{'='*70}")

log = Log()


# ============================================================
# SETUP
# ============================================================
def setup():
    log.section("Setting up output folder")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    log.ok(f"Folder: {OUTPUT_DIR}")
    expected = {
        'Table1_csv':     STATS_DIR / "Table1_regional_climate_summary.csv",
        'Table2_csv':     STATS_DIR / "Table2_trend_statistics_ModifiedMK.csv",
        'Table3_csv':     STATS_DIR / "TableB4_Decadal_Shift_cumulative_STRICTER_FDR.csv",
        'TableS1_csv':    STATS_DIR / "TableS1_SPEI_validation_diagnostics.csv",
        'TableS2_csv':    STATS_DIR / "TableS2_MK_comparison_with_lag1.csv",
        'TableS3_csv':    STATS_DIR / "TableS3_robustness_area_weighted.csv",
        'TableS3b_csv':   STATS_DIR / "TableS3b_continental_degradation_by_pathway.csv",
        'aridity_class':  FIGS_DIR  / "aridity_classes_UNEP_1985_2022.tif",
        'aridity_PoverPET': FIGS_DIR / "aridity_PoverPET_1985_2022.tif",
        'figure1_drought_trends': FIGS_DIR / "Figure1_drought_trends.png",
        'figure2_AED':    FIGS_DIR  / "Figure2_AED_and_frequency.png",
        'figure3_ts':     FIGS_DIR  / "Figure3_regional_timeseries.png",
        'figure4_intervals': FIGS_DIR / "Figure4_SPEI_intervals.png",
        'figure6_heatmap': FIGS_DIR / "Figure6_decadal_heatmap.png",
        'figureB1_data_lagged':     MODULEB_DIR / "FigureB1_data_lagged.csv",
        'figureB1_data_cumulative': MODULEB_DIR / "FigureB1_data_cumulative.csv",
        'figureB2_data':            MODULEB_DIR / "FigureB2_data.csv",
        'figureB3_data':            MODULEB_DIR / "FigureB3_data.csv",
        'figureB4_data':            MODULEB_DIR / "FigureB4_data.csv",
        'figureB5_data_RSI':        MODULEB_DIR / "FigureB5_data_RSI.csv",
        'figureB5_data_DRA':        MODULEB_DIR / "FigureB5_data_DRA.csv",
        'figureB1_png': MODULEB_DIR / "FigureB1_attribution_heatmap.png",
        'figureB2_png': MODULEB_DIR / "FigureB2_recovery_suppression.png",
        'figureB3_png': MODULEB_DIR / "FigureB3_category_comparison.png",
        'figureB4_png': MODULEB_DIR / "FigureB4_decadal_shift.png",
        'figureB5_png': MODULEB_DIR / "FigureB5_timescale_sensitivity.png",
        'shapefile':    SHAPEFILE,
    }
    found = sum(1 for p in expected.values() if p.exists())
    log.info(f"Found {found}/{len(expected)} input files.")
    return expected


def save_fig(fig, name):
    path = OUTPUT_DIR / f"{name}.png"
    fig.savefig(path, dpi=DPI, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ============================================================
# FIGURE 1, 2, 3 — UNCHANGED FROM v1 (they worked)
# ============================================================
def figure_1(inputs):
    log.section("FIGURE 1 — Study area + aridity")
    if not HAS_RASTERIO or not HAS_GEOPANDAS:
        log.skip("Need rasterio + geopandas"); return
    arid_class_path = inputs['aridity_class']
    arid_cont_path  = inputs['aridity_PoverPET']
    shp_path        = inputs['shapefile']
    if not arid_class_path.exists() or not shp_path.exists():
        log.skip("Aridity raster or shapefile missing"); return
    try:
        regions = gpd.read_file(shp_path)
        if regions.crs is None:
            regions = regions.set_crs("EPSG:4326")
        regions = regions.to_crs("EPSG:4326")
        with rasterio.open(arid_class_path) as src:
            arid_class = src.read(1, masked=True)
            extent = [src.bounds.left, src.bounds.right,
                      src.bounds.bottom, src.bounds.top]
        fig, axes = plt.subplots(1, 2, figsize=(13, 6.5))
        for ax in axes:
            ax.set_aspect('equal')
            ax.set_xlabel('Longitude'); ax.set_ylabel('Latitude')
        ax = axes[0]
        for code in REGION_ORDER:
            sel = regions[regions['LAB'] == code] if 'LAB' in regions.columns else regions.iloc[0:0]
            if len(sel) > 0:
                sel.plot(ax=ax, facecolor=REGION_COLORS[code], edgecolor='k',
                         linewidth=0.5, alpha=0.6,
                         label=f"{code} — {REGION_NAMES[code]}")
        ax.set_title("(a) IPCC AR5 reference regions")
        ax.legend(loc='lower left', fontsize=8)
        ax.set_xlim(-20, 55); ax.set_ylim(-37, 38)
        ax = axes[1]
        class_colors = ['#a50f15','#fcae91','#f7f4c0','#a1d99b','#41ab5d','#1f6628']
        class_labels = ['Hyper-arid','Arid','Semi-arid','Dry sub-humid','Humid','Very humid']
        bounds = [0.5,1.5,2.5,3.5,4.5,5.5,6.5]
        cmap = ListedColormap(class_colors); norm = BoundaryNorm(bounds, cmap.N)
        ax.imshow(arid_class, extent=extent, origin='upper', cmap=cmap, norm=norm,
                  interpolation='nearest')
        regions.boundary.plot(ax=ax, edgecolor='k', linewidth=0.6)
        ax.set_title("(b) UNEP aridity classes (mean 1985–2022)")
        ax.set_xlim(-20, 55); ax.set_ylim(-37, 38)
        patches = [Patch(facecolor=c, edgecolor='k', label=l)
                   for c, l in zip(class_colors, class_labels)]
        ax.legend(handles=patches, loc='lower left', fontsize=7, ncol=2)
        plt.tight_layout()
        out = save_fig(fig, "Figure1_study_area")
        log.ok(f"Figure 1 → {out.name}")
        shutil.copy2(arid_class_path,
                     OUTPUT_DIR / "Figure1_panel_b_UNEP_classes.tif")
        log.ok("Source TIF copied (UNEP classes)")
        if arid_cont_path.exists():
            shutil.copy2(arid_cont_path,
                         OUTPUT_DIR / "Figure1_panel_b_PoverPET.tif")
            log.ok("Source TIF copied (P/PET continuous)")
    except Exception as e:
        log.fail(f"Figure 1 failed: {e}")


def figure_2(inputs):
    log.section("FIGURE 2 — AED + drought frequency gap")
    src = inputs['figure2_AED']
    if not src.exists():
        log.skip("Missing"); return
    img = plt.imread(src)
    h, w = img.shape[:2]
    fig, ax = plt.subplots(figsize=(w/100, h/100))
    ax.imshow(img); ax.axis('off')
    plt.tight_layout(pad=0)
    out = save_fig(fig, "Figure2_climate_baseline")
    log.ok(f"Figure 2 → {out.name}")


def figure_3(inputs):
    log.section("FIGURE 3 — Drought intensification (4-panel)")
    src_trends = inputs['figure1_drought_trends']
    src_ts     = inputs['figure3_ts']
    src_decade = inputs['figure6_heatmap']
    for name, p in [('trends', src_trends), ('time series', src_ts),
                    ('decadal heatmap', src_decade)]:
        if not p.exists():
            log.skip(f"Missing: {name}"); return
    img_trends = plt.imread(src_trends)
    img_ts     = plt.imread(src_ts)
    img_decade = plt.imread(src_decade)
    h, w = img_trends.shape[:2]
    half_h, half_w = h // 2, w // 2
    panel_a = img_trends[:half_h,   :half_w]
    panel_b = img_trends[half_h:,   half_w:]
    fig, axes = plt.subplots(2, 2, figsize=(14, 11))
    axes[0, 0].imshow(panel_a); axes[0, 0].set_title("(a) SPEI-12 trend", loc='left'); axes[0, 0].axis('off')
    axes[0, 1].imshow(panel_b); axes[0, 1].set_title("(b) PET trend", loc='left'); axes[0, 1].axis('off')
    axes[1, 0].imshow(img_ts); axes[1, 0].set_title("(c) Regional-mean SPEI-12 time series", loc='left'); axes[1, 0].axis('off')
    axes[1, 1].imshow(img_decade); axes[1, 1].set_title("(d) Decadal-mean SPEI-12 by region", loc='left'); axes[1, 1].axis('off')
    plt.tight_layout(pad=1.0)
    out = save_fig(fig, "Figure3_drought_intensification")
    log.ok(f"Figure 3 → {out.name}")


# ============================================================
# FIGURE 4 — DEFENSIVE CSV-BASED RENDER
# ============================================================
def figure_4(inputs):
    log.section("FIGURE 4 — Pathway × region attribution heatmap (CSV)")
    src_lag = inputs['figureB1_data_lagged']
    src_cum = inputs['figureB1_data_cumulative']
    if not src_lag.exists() or not src_cum.exists():
        log.skip("CSVs missing"); return

    df_lag = pd.read_csv(src_lag)
    df_cum = pd.read_csv(src_cum)
    print_columns("FigureB1_data_lagged", df_lag)
    print_columns("FigureB1_data_cumulative", df_cum)

    def make_pivot(df):
        path_col = find_col(df, ['pathway', 'transition', 'path', 'trans'], 'pathway')
        reg_col  = find_col(df, ['region', 'reg', 'lab'],                   'region')
        d_col    = find_col(df, ['cohens_d', 'mean_d', 'cohen', 'd_mean', 'value', 'd'], 'd')
        p_col    = find_col(df, ['p_value', 'pval', 'p_val', 'pvalue'],     'p')
        if not all([path_col, reg_col, d_col]):
            return None, None
        # Aggregate in case there are duplicates
        agg = df.groupby([path_col, reg_col]).agg({d_col: 'mean'}).reset_index()
        pivot_d = agg.pivot(index=path_col, columns=reg_col, values=d_col)
        pivot_d = pivot_d.reindex(columns=[c for c in REGION_ORDER if c in pivot_d.columns])
        pivot_p = None
        if p_col:
            agg_p = df.groupby([path_col, reg_col]).agg({p_col: 'mean'}).reset_index()
            pivot_p = agg_p.pivot(index=path_col, columns=reg_col, values=p_col)
            pivot_p = pivot_p.reindex(columns=[c for c in REGION_ORDER if c in pivot_p.columns])
        # Reorder pathways
        pathway_order = ['FST_SHR','SHR_GRS','FST_CRP','GRS_BAL',
                         'SHR_FST','GRS_SHR','BAL_GRS',
                         'AGEXPANSION','CRP_ABANDONMENT']
        present = [p for p in pathway_order if p in pivot_d.index]
        # Add any other pathways at the bottom
        others = [p for p in pivot_d.index if p not in present]
        pivot_d = pivot_d.reindex(present + others)
        if pivot_p is not None:
            pivot_p = pivot_p.reindex(present + others)
        return pivot_d, pivot_p

    pivot_d_lag, pivot_p_lag = make_pivot(df_lag)
    pivot_d_cum, pivot_p_cum = make_pivot(df_cum)
    if pivot_d_lag is None or pivot_d_cum is None:
        log.fail("Could not build pivot tables — required columns missing")
        return

    fig, axes = plt.subplots(1, 2, figsize=(14, 7))
    cmap = mpl.cm.RdBu_r
    norm = mpl.colors.TwoSlopeNorm(vcenter=0, vmin=-0.6, vmax=0.6)
    im = None
    for ax, pivot_d, pivot_p, title in [
        (axes[0], pivot_d_lag, pivot_p_lag, "Lagged window (primary)"),
        (axes[1], pivot_d_cum, pivot_p_cum, "Cumulative window (robustness)"),
    ]:
        im = ax.imshow(pivot_d.values, cmap=cmap, norm=norm, aspect='auto')
        ax.set_xticks(range(len(pivot_d.columns)))
        ax.set_xticklabels(pivot_d.columns, fontsize=10)
        ax.set_yticks(range(len(pivot_d.index)))
        ax.set_yticklabels(pivot_d.index, fontsize=9)
        ax.set_title(title)
        for i in range(pivot_d.shape[0]):
            for j in range(pivot_d.shape[1]):
                v = pivot_d.iat[i, j]
                if pd.isna(v): continue
                star = ''
                if pivot_p is not None and j < pivot_p.shape[1] and i < pivot_p.shape[0]:
                    pv = pivot_p.iat[i, j]
                    if not pd.isna(pv) and pv < 0.05: star = '*'
                ax.text(j, i, f"{v:+.2f}{star}", ha='center', va='center',
                        fontsize=8.5, color='black' if abs(v) < 0.35 else 'white')
        ax.set_xlabel("Region")
    fig.colorbar(im, ax=axes, shrink=0.7, pad=0.02,
                 label="Cohen's d (transition vs stable)")
    fig.suptitle("Attribution heatmap — pathway × region\n* one-sample t-test p < 0.05",
                 fontsize=11, y=1.02)
    out = save_fig(fig, "Figure4_attribution_heatmap")
    log.ok(f"Figure 4 → {out.name}")


# ============================================================
# FIGURE 5 — DEFENSIVE CSV-BASED RENDER
# ============================================================
def figure_5(inputs):
    log.section("FIGURE 5 — RSI + timescale sensitivity")
    src_rsi    = inputs['figureB2_data']
    src_ts_rsi = inputs['figureB5_data_RSI']
    src_ts_dra = inputs['figureB5_data_DRA']
    for p in [src_rsi, src_ts_rsi, src_ts_dra]:
        if not p.exists():
            log.skip(f"Missing: {p}"); return

    df_rsi    = pd.read_csv(src_rsi);    print_columns("FigureB2_data", df_rsi)
    df_ts_rsi = pd.read_csv(src_ts_rsi); print_columns("FigureB5_data_RSI", df_ts_rsi)
    df_ts_dra = pd.read_csv(src_ts_dra); print_columns("FigureB5_data_DRA", df_ts_dra)

    reg_col = find_col(df_rsi, ['region', 'reg', 'lab'], 'region (B2)')
    win_col = find_col(df_rsi, ['window', 'win'],        'window (B2)')
    val_col = find_col(df_rsi, ['rsi', 'mean_d', 'cohens_d', 'value', 'd'], 'value (B2)')
    p_col   = find_col(df_rsi, ['p_value', 'pval', 'p_val'], 'p (B2)')

    if not all([reg_col, val_col]):
        log.fail("Required B2 columns missing")
        return

    fig = plt.figure(figsize=(16, 8))
    gs = fig.add_gridspec(2, 4, width_ratios=[1.4,1,1,1], hspace=0.4, wspace=0.4)

    # Panel (a) — RSI bars
    ax_a = fig.add_subplot(gs[:, 0])

    if win_col:
        lag = df_rsi[df_rsi[win_col].astype(str).str.lower().str.contains('lag', na=False)]
        cum = df_rsi[df_rsi[win_col].astype(str).str.lower().str.contains('cum', na=False)]
    else:
        lag = df_rsi; cum = df_rsi.iloc[0:0]

    def get_region_values(df_sub):
        out = []
        out_p = []
        for r in REGION_ORDER:
            row = df_sub[df_sub[reg_col] == r]
            v = row[val_col].iloc[0] if len(row) else np.nan
            pv = row[p_col].iloc[0] if (p_col and len(row)) else np.nan
            out.append(v); out_p.append(pv)
        return out, out_p

    lag_vals, lag_p = get_region_values(lag)
    cum_vals, cum_p = get_region_values(cum)

    x = np.arange(len(REGION_ORDER))
    w = 0.36
    ax_a.bar(x - w/2, lag_vals, w, color='#1f77b4', edgecolor='k',
             linewidth=0.5, label='Lagged window (primary)')
    if any(not pd.isna(v) for v in cum_vals):
        ax_a.bar(x + w/2, cum_vals, w, color='#9467bd', edgecolor='k',
                 linewidth=0.5, label='Cumulative window')

    for i, (v, pv) in enumerate(zip(lag_vals, lag_p)):
        if not pd.isna(v):
            ax_a.text(i - w/2, v + (0.02 if v >= 0 else -0.02), f"{v:+.2f}",
                      ha='center', va='bottom' if v >= 0 else 'top', fontsize=8)
            if not pd.isna(pv) and pv < 0.05:
                ax_a.text(i - w/2, v + (0.05 if v >= 0 else -0.05), '*',
                          ha='center', fontsize=14, fontweight='bold')
    for i, (v, pv) in enumerate(zip(cum_vals, cum_p)):
        if not pd.isna(v):
            ax_a.text(i + w/2, v + (0.02 if v >= 0 else -0.02), f"{v:+.2f}",
                      ha='center', va='bottom' if v >= 0 else 'top', fontsize=8)
            if not pd.isna(pv) and pv < 0.05:
                ax_a.text(i + w/2, v + (0.05 if v >= 0 else -0.05), '*',
                          ha='center', fontsize=14, fontweight='bold')

    ax_a.axhline(0, color='k', linewidth=0.8)
    ax_a.set_xticks(x)
    ax_a.set_xticklabels([REGION_NAMES[r] for r in REGION_ORDER], rotation=20, ha='right')
    ax_a.set_ylabel("Recovery Suppression Index (Cohen's d)")
    ax_a.set_title("(a) Recovery Suppression Index by region (SPEI-12)", loc='left')
    ax_a.legend(loc='lower right', fontsize=9)
    ax_a.text(0.02, 0.98, '* p < 0.05', transform=ax_a.transAxes,
              va='top', fontsize=9,
              bbox=dict(boxstyle='round', facecolor='white', edgecolor='gray'))

    # Panel (b) — timescale grids
    def render_ts_panel(ax, df, value_col_keywords, marker, title, ylabel):
        ts_col  = find_col(df, ['timescale', 'scale', 'months', 'period'], 'timescale')
        reg_col_x = find_col(df, ['region', 'reg', 'lab'], 'region')
        v_col   = find_col(df, value_col_keywords, 'value')
        p_col_x = find_col(df, ['p_value', 'pval', 'p_val'], 'p')
        if not all([ts_col, reg_col_x, v_col]):
            ax.text(0.5, 0.5, "data unavailable", ha='center', transform=ax.transAxes)
            return
        for r in REGION_ORDER:
            sub = df[df[reg_col_x] == r].sort_values(ts_col)
            if len(sub) == 0: continue
            ax.plot(sub[ts_col], sub[v_col], f'{marker}-',
                    color=REGION_COLORS[r], markersize=6, label=REGION_NAMES[r])
            if p_col_x:
                for _, row in sub.iterrows():
                    if not pd.isna(row[p_col_x]) and row[p_col_x] < 0.05:
                        ax.text(row[ts_col], row[v_col] - 0.02, '*',
                                ha='center', fontsize=12, fontweight='bold',
                                color=REGION_COLORS[r])
        ax.axhline(0, color='gray', linewidth=0.5, linestyle='--')
        ax.set_xlabel('SPEI timescale (months)')
        ax.set_ylabel(ylabel)
        ax.set_title(title, loc='left', fontsize=10)

    # Split RSI and DRA dataframes by window
    win_col_rsi = find_col(df_ts_rsi, ['window', 'win'], 'window (B5 RSI)')
    win_col_dra = find_col(df_ts_dra, ['window', 'win'], 'window (B5 DRA)')

    def split_by_window(df, win_col):
        if win_col is None:
            return df, df.iloc[0:0]
        lag = df[df[win_col].astype(str).str.lower().str.contains('lag', na=False)]
        cum = df[df[win_col].astype(str).str.lower().str.contains('cum', na=False)]
        return lag, cum

    rsi_lag, rsi_cum = split_by_window(df_ts_rsi, win_col_rsi)
    dra_lag, dra_cum = split_by_window(df_ts_dra, win_col_dra)

    ax = fig.add_subplot(gs[0, 1])
    render_ts_panel(ax, rsi_lag, ['rsi','mean_d','value','cohens_d','d'],
                    'o', "(b1) RSI — lagged", "RSI (Cohen's d)")
    ax = fig.add_subplot(gs[0, 2])
    render_ts_panel(ax, rsi_cum, ['rsi','mean_d','value','cohens_d','d'],
                    'o', "(b2) RSI — cumulative", "RSI (Cohen's d)")
    ax = fig.add_subplot(gs[1, 1])
    render_ts_panel(ax, dra_lag, ['dra','mean_d','value','d_diff','d'],
                    's', "(b3) DRA — lagged", "DRA (Cohen's d diff.)")
    ax = fig.add_subplot(gs[1, 2])
    render_ts_panel(ax, dra_cum, ['dra','mean_d','value','d_diff','d'],
                    's', "(b4) DRA — cumulative", "DRA (Cohen's d diff.)")

    # Region legend below
    handles = [plt.Line2D([0], [0], marker='o', color=REGION_COLORS[r],
                          label=REGION_NAMES[r], markersize=7, linewidth=0)
               for r in REGION_ORDER]
    fig.legend(handles=handles, loc='lower center', ncol=5,
               bbox_to_anchor=(0.6, -0.02), frameon=True, fontsize=9)
    fig.suptitle("Figure 5 — RSI and timescale sensitivity", y=1.00, fontsize=12)
    out = save_fig(fig, "Figure5_RSI_timescale")
    log.ok(f"Figure 5 → {out.name}")


# ============================================================
# FIGURE 6 — DEFENSIVE CSV-BASED RENDER
# ============================================================
def figure_6(inputs):
    log.section("FIGURE 6 — Decadal shift in attribution")
    src = inputs['figureB4_data']
    if not src.exists():
        log.skip("Missing"); return
    df = pd.read_csv(src); print_columns("FigureB4_data", df)

    reg_col = find_col(df, ['region','reg','lab'],           'region')
    cat_col = find_col(df, ['category','cat'],               'category')
    per_col = find_col(df, ['period','window','epoch','era'], 'period')
    val_col = find_col(df, ['mean_d','cohens_d','d_mean','value','d'], 'value')
    surv_col= find_col(df, ['fdr','survives','sig'],         'fdr')

    if not all([reg_col, cat_col, val_col]):
        log.fail("Required B4 columns missing")
        return

    categories = ['Degradation','Recovery','Agricultural']
    cat_short  = {'Degradation':'Deg','Recovery':'Rec','Agricultural':'Agr'}
    color_early= {'Degradation':'#f4cccc','Recovery':'#d9ead3','Agricultural':'#fce5cd'}
    color_late = {'Degradation':'#cc0000','Recovery':'#38761d','Agricultural':'#e69138'}

    fig, ax = plt.subplots(figsize=(14, 6.5))
    cat_width = 0.85 / 3; period_width = cat_width / 2.2

    for ri, region in enumerate(REGION_ORDER):
        for ci, cat in enumerate(categories):
            sub = df[(df[reg_col] == region) &
                     (df[cat_col].astype(str).str.lower() == cat.lower())]
            if len(sub) == 0: continue
            if per_col:
                early = sub[sub[per_col].astype(str).str.lower().str.contains('early', na=False)]
                late  = sub[sub[per_col].astype(str).str.lower().str.contains('late',  na=False)]
            else:
                early = sub.iloc[:1]; late = sub.iloc[1:2]
            xpos_c = ri + (ci - 1) * cat_width
            if len(early):
                v = early[val_col].iloc[0]
                ax.bar(xpos_c - period_width/2, v, period_width*0.9,
                       color=color_early[cat], edgecolor='k', linewidth=0.5)
            if len(late):
                v = late[val_col].iloc[0]
                ax.bar(xpos_c + period_width/2, v, period_width*0.9,
                       color=color_late[cat], edgecolor='k', linewidth=0.5)
            if surv_col and any(sub[surv_col] == True):
                ymax = sub[val_col].max()
                ax.text(xpos_c, ymax + 0.10, '*', ha='center',
                        fontsize=16, fontweight='bold')

    ax.axhline(0, color='k', linewidth=0.8)
    ax.set_xticks(range(len(REGION_ORDER)))
    ax.set_xticklabels([REGION_NAMES[r] for r in REGION_ORDER])
    for ri in range(len(REGION_ORDER)):
        for ci, cat in enumerate(categories):
            ax.text(ri + (ci - 1) * cat_width, -0.85, cat_short[cat],
                    ha='center', fontsize=8, color='gray')
    ax.set_ylabel("Cohen's d"); ax.set_ylim(-1.0, 0.8)
    ax.set_title("Decadal shift in attribution: early (1990–2004) vs late (2005–2022)\n"
                 "Cumulative SPEI-12, Welch's t-test with Benjamini-Hochberg FDR\n"
                 "* survives FDR (α = 0.05)")

    handles = []
    for cat in categories:
        handles.append(Patch(facecolor=color_early[cat], edgecolor='k',
                             label=f'{cat} early'))
        handles.append(Patch(facecolor=color_late[cat], edgecolor='k',
                             label=f'{cat} late'))
    ax.legend(handles=handles, loc='upper right', fontsize=8, ncol=2)
    out = save_fig(fig, "Figure6_decadal_shift")
    log.ok(f"Figure 6 → {out.name}")


# ============================================================
# FIGURE 7 — DEFENSIVE CSV-BASED RENDER
# ============================================================
def figure_7(inputs):
    log.section("FIGURE 7 — Cohen's d distributions")
    src = inputs['figureB3_data']
    if not src.exists():
        log.skip("Missing"); return
    df = pd.read_csv(src); print_columns("FigureB3_data", df)

    win_col = find_col(df, ['window','win'],              'window')
    cat_col = find_col(df, ['category','cat'],            'category')
    reg_col = find_col(df, ['region','reg','lab'],        'region')
    val_col = find_col(df, ['cohens_d','mean_d','value','d'], 'd')

    if not all([reg_col, cat_col, val_col]):
        log.fail("Required B3 columns missing")
        return

    categories = ['Degradation','Recovery','Agricultural']
    cat_color = {'Degradation':'#d62728','Recovery':'#2ca02c','Agricultural':'#ff7f0e'}
    cat_short = {'Degradation':'Deg','Recovery':'Rec','Agricultural':'Agr'}

    if win_col:
        windows = [('lagged', df[df[win_col].astype(str).str.lower().str.contains('lag', na=False)]),
                   ('cumulative', df[df[win_col].astype(str).str.lower().str.contains('cum', na=False)])]
    else:
        windows = [('all data', df)]

    n_panels = len(windows)
    fig, axes = plt.subplots(1, n_panels, figsize=(8*n_panels, 7), sharey=True)
    if n_panels == 1: axes = [axes]

    for ax, (win_label, sub_df) in zip(axes, windows):
        positions, labels, colors, boxes = [], [], [], []
        pos = 0
        for r in REGION_ORDER:
            for cat in categories:
                vals = sub_df[(sub_df[reg_col] == r) &
                              (sub_df[cat_col].astype(str).str.lower() == cat.lower())][val_col].dropna().values
                if len(vals) > 0:
                    positions.append(pos); boxes.append(vals)
                    colors.append(cat_color[cat])
                    labels.append(f"{r}\n{cat_short[cat]}")
                pos += 1
            pos += 0.5
        if boxes:
            bp = ax.boxplot(boxes, positions=positions, widths=0.7,
                            patch_artist=True, showmeans=True, showfliers=False,
                            meanprops=dict(marker='D', markerfacecolor='white',
                                          markeredgecolor='k', markersize=5),
                            medianprops=dict(color='k', linewidth=1.5))
            for patch, c in zip(bp['boxes'], colors):
                patch.set_facecolor(c); patch.set_alpha(0.7); patch.set_edgecolor('k')
            ax.set_xticks(positions); ax.set_xticklabels(labels, fontsize=7)
        ax.axhline(0, color='k', linewidth=0.6, linestyle='--')
        ax.set_ylim(-1, 1); ax.set_title(f"{win_label.capitalize()} window")
        if ax is axes[0]: ax.set_ylabel("Cohen's d")
    handles = [Patch(facecolor=c, alpha=0.7, edgecolor='k', label=cat)
               for cat, c in cat_color.items()]
    fig.legend(handles=handles, loc='upper center', ncol=3,
               bbox_to_anchor=(0.5, 1.02), frameon=True)
    fig.suptitle("Distribution of Cohen's d per category × region (SPEI-12)",
                 y=1.06, fontsize=12)
    out = save_fig(fig, "Figure7_distributions")
    log.ok(f"Figure 7 → {out.name}")


# ============================================================
# SUPPLEMENTARY FIGURES
# ============================================================
def figure_s1(inputs):
    log.section("SUP. FIG. S1 — Per-interval SPEI-12")
    src = inputs['figure4_intervals']
    if not src.exists():
        log.skip("Missing"); return
    img = plt.imread(src); h, w = img.shape[:2]
    fig, ax = plt.subplots(figsize=(w/100, h/100))
    ax.imshow(img); ax.axis('off'); plt.tight_layout(pad=0)
    out = save_fig(fig, "FigureS1_SPEI_intervals")
    log.ok(f"Sup. Fig. S1 → {out.name}")


def figure_s2(inputs):
    log.section("SUP. FIG. S2 — Full timescale grid")
    src = inputs['figureB5_png']
    if not src.exists():
        log.skip("Missing"); return
    img = plt.imread(src); h, w = img.shape[:2]
    fig, ax = plt.subplots(figsize=(w/100, h/100))
    ax.imshow(img); ax.axis('off'); plt.tight_layout(pad=0)
    out = save_fig(fig, "FigureS2_timescale_grid")
    log.ok(f"Sup. Fig. S2 → {out.name}")


# ============================================================
# TABLES
# ============================================================
def write_word_table(df, headers, output_path, title=None, footnote=None):
    if not HAS_DOCX: return
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'; style.font.size = Pt(10)
    if title:
        p = doc.add_paragraph(); run = p.add_run(title)
        run.bold = True; run.font.size = Pt(11)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for para in cell.paragraphs:
            for r in para.runs: r.bold = True; r.font.size = Pt(9)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, h in enumerate(headers):
            cells[i].text = str(row.get(h, ''))
            for para in cells[i].paragraphs:
                for r in para.runs: r.font.size = Pt(9)
    if footnote:
        p = doc.add_paragraph(); r = p.add_run(footnote)
        r.italic = True; r.font.size = Pt(8)
    doc.save(output_path)


def table_1(inputs):
    log.section("TABLE 1 — Regional climate")
    src = inputs['Table1_csv']
    if not src.exists(): log.skip("Missing"); return
    df = pd.read_csv(src)
    cols = ['region','region_name','mean_annual_precip_mm','mean_annual_PET_mm',
            'mean_SPEI12','mean_SPI12','aed_contribution_pct',
            'pct_SPEI_drought_below_m1','pct_SPI_drought_below_m1']
    df_out = df[[c for c in cols if c in df.columns]].copy()
    rounding = {'mean_annual_precip_mm':0,'mean_annual_PET_mm':0,
                'mean_SPEI12':2,'mean_SPI12':2,'aed_contribution_pct':1,
                'pct_SPEI_drought_below_m1':1,'pct_SPI_drought_below_m1':1}
    for c, n in rounding.items():
        if c in df_out.columns: df_out[c] = df_out[c].round(n)
    df_out.to_csv(OUTPUT_DIR / "Table1_regional_climate_summary.csv", index=False)
    log.ok("Table 1 CSV → Table1_regional_climate_summary.csv")
    if HAS_DOCX:
        rename = {'region':'Region','region_name':'Region name',
                  'mean_annual_precip_mm':'Precip (mm/yr)','mean_annual_PET_mm':'PET (mm/yr)',
                  'mean_SPEI12':'Mean SPEI-12','mean_SPI12':'Mean SPI-12',
                  'aed_contribution_pct':'AED contrib. (%)',
                  'pct_SPEI_drought_below_m1':'SPEI < −1 (%)',
                  'pct_SPI_drought_below_m1':'SPI < −1 (%)'}
        df_word = df_out.rename(columns=rename)
        write_word_table(df_word.astype(str), list(df_word.columns),
                         OUTPUT_DIR / "Table1_regional_climate_summary.docx",
                         title="Table 1. Regional climate baseline, 1985–2022.",
                         footnote="SPEI-12 / SPI-12 are means across all months 1985–2022. "
                                  "Drought-exposure columns: % pixel-months with index < −1. "
                                  "AED contribution = (SPEI severity − SPI severity)/SPEI severity, %.")
        log.ok("Table 1 Word → Table1_regional_climate_summary.docx")


def table_2(inputs):
    log.section("TABLE 2 — Mann-Kendall trends")
    src = inputs['Table2_csv']
    if not src.exists(): log.skip("Missing"); return
    df = pd.read_csv(src)
    cols = ['region','region_name','SPEI12_mean_slope','SPEI12_pct_drying_sig',
            'SPEI12_pct_wetting_sig','SPI12_mean_slope','SPI12_pct_drying_sig',
            'SPI12_pct_wetting_sig','precip_mean_slope','precip_pct_drying_sig',
            'precip_pct_wetting_sig','PET_mean_slope','PET_pct_significant']
    df_out = df[[c for c in cols if c in df.columns]].copy()
    rounding = {'SPEI12_mean_slope':4,'SPEI12_pct_drying_sig':1,'SPEI12_pct_wetting_sig':1,
                'SPI12_mean_slope':4,'SPI12_pct_drying_sig':1,'SPI12_pct_wetting_sig':1,
                'precip_mean_slope':2,'precip_pct_drying_sig':1,'precip_pct_wetting_sig':1,
                'PET_mean_slope':2,'PET_pct_significant':1}
    for c, n in rounding.items():
        if c in df_out.columns: df_out[c] = df_out[c].round(n)
    df_out.to_csv(OUTPUT_DIR / "Table2_trend_statistics_ModifiedMK.csv", index=False)
    log.ok("Table 2 CSV → Table2_trend_statistics_ModifiedMK.csv")
    if HAS_DOCX:
        rename = {'region':'Region','region_name':'Region name',
                  'SPEI12_mean_slope':'SPEI-12 slope','SPEI12_pct_drying_sig':'SPEI-12 % drying',
                  'SPEI12_pct_wetting_sig':'SPEI-12 % wetting',
                  'SPI12_mean_slope':'SPI-12 slope','SPI12_pct_drying_sig':'SPI-12 % drying',
                  'SPI12_pct_wetting_sig':'SPI-12 % wetting',
                  'precip_mean_slope':'Precip slope (mm/yr)','precip_pct_drying_sig':'Precip % drying',
                  'precip_pct_wetting_sig':'Precip % wetting',
                  'PET_mean_slope':'PET slope (mm/yr)','PET_pct_significant':'PET % sig.'}
        df_word = df_out.rename(columns=rename)
        write_word_table(df_word.astype(str), list(df_word.columns),
                         OUTPUT_DIR / "Table2_trend_statistics_ModifiedMK.docx",
                         title="Table 2. Modified Mann-Kendall trends, 1985–2022.",
                         footnote="Hamed-Rao variance correction. Theil-Sen slopes. "
                                  "Percent = pixels with significant trend (α=0.05) by direction.")
        log.ok("Table 2 Word → Table2_trend_statistics_ModifiedMK.docx")


def table_3(inputs):
    log.section("TABLE 3 — FDR-significant decadal shifts")
    src = inputs['Table3_csv']
    if not src.exists(): log.skip("Missing"); return
    df = pd.read_csv(src)
    sig_col = next((c for c in df.columns if 'fdr' in c.lower() and 'sig' in c.lower()), None)
    if sig_col is None:
        log.skip("No FDR flag column"); return
    df_sig = df[df[sig_col] == True].copy()
    sort_col = 'q_value_bh' if 'q_value_bh' in df.columns else None
    if sort_col: df_sig = df_sig.sort_values(sort_col)
    df_sig.to_csv(OUTPUT_DIR / "Table3_FDR_significant_decadal_shifts.csv", index=False)
    log.ok(f"Table 3 CSV → Table3_FDR_significant_decadal_shifts.csv ({len(df_sig)} rows)")
    if HAS_DOCX:
        write_word_table(df_sig.astype(str), list(df_sig.columns),
                         OUTPUT_DIR / "Table3_FDR_significant_decadal_shifts.docx",
                         title="Table 3. FDR-significant decadal shifts (cumulative SPEI-12).",
                         footnote="Five region × category combinations surviving BH-FDR (α=0.05) "
                                  "from 15 Welch's t-tests. Δd = late (2005-2022) − early (1990-2004) "
                                  "period mean. All negative: late-period transitions in wetter conditions.")
        log.ok("Table 3 Word → Table3_FDR_significant_decadal_shifts.docx")


def supplementary_tables(inputs):
    log.section("SUPPLEMENTARY TABLES — copy with publication names")
    pairs = [
        (inputs['TableS1_csv'],  "TableS1_SPEI_validation_diagnostics.csv"),
        (inputs['TableS2_csv'],  "TableS2_MK_comparison_with_lag1.csv"),
        (inputs['TableS3_csv'],  "TableS3_robustness_area_weighted.csv"),
        (inputs['TableS3b_csv'], "TableS3b_continental_degradation_by_pathway.csv"),
        (inputs['Table3_csv'],   "TableS4_Decadal_Shift_full_panel_FDR.csv"),
    ]
    for src, dst_name in pairs:
        if src.exists():
            shutil.copy2(src, OUTPUT_DIR / dst_name)
            log.ok(f"Sup. Table → {dst_name}")
        else:
            log.skip(f"Missing: {src.name}")
    src_extra = STATS_DIR / "TableS3b_continental_degradation_by_group.csv"
    if src_extra.exists():
        shutil.copy2(src_extra,
                     OUTPUT_DIR / "TableS3b_continental_degradation_by_group.csv")
        log.ok("Sup. Table → TableS3b_continental_degradation_by_group.csv")


# ============================================================
# MAIN
# ============================================================
def main():
    print(f"\n{'='*70}\nPublication figures & tables generator (v2)\n"
          f"Output: {OUTPUT_DIR}\nDPI: {DPI}\n{'='*70}")
    inputs = setup()
    figure_1(inputs); figure_2(inputs); figure_3(inputs)
    figure_4(inputs); figure_5(inputs); figure_6(inputs); figure_7(inputs)
    figure_s1(inputs); figure_s2(inputs)
    table_1(inputs); table_2(inputs); table_3(inputs); supplementary_tables(inputs)
    log.summary()
    with open(OUTPUT_DIR / "_MANIFEST.txt", 'w') as f:
        f.write("Publication figures & tables — manifest\n" + "="*60 + "\n\n")
        for s, m in log.entries: f.write(f"[{s}] {m}\n")
    print(f"Manifest: {OUTPUT_DIR / '_MANIFEST.txt'}")


if __name__ == "__main__":
    main()
